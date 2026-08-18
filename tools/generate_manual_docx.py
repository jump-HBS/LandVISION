# -*- coding: utf-8 -*-
"""
《项目说明书.docx》生成器 —— 细化到项目内每个文件的职责说明。

用法：
    cd C:\\landvision-project
    venv\\Scripts\\python.exe tools\\generate_manual_docx.py

产物：项目根目录《项目说明书.docx》。
说明：本脚本遍历项目目录（排除 venv/node_modules/.git/dist/缓存/日志），
     每个文件在 DESCRIPTIONS 中必须有对应说明；缺失时打印警告并写入占位文本。
"""
import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT_DOCX = os.path.join(ROOT, "项目说明书.docx")

EXCLUDE_DIRS = {"venv", "node_modules", ".git", "dist", "__pycache__",
                ".pytest_cache", "logs", ".vscode"}
EXCLUDE_FILES = {"项目说明书.docx"}
# ---------------------------------------------------------------------------
# 逐文件说明（键为相对路径，统一用 / 分隔）
# ---------------------------------------------------------------------------
DESCRIPTIONS = {
    # ---------- 根目录 ----------
    "README.md": "项目总览与快速开始：一句话定位、双模式启动命令、目录结构、六层架构表、接口总览、页面总览、学习路线。新同学看这一份就够了。",
    "项目实现框架.md": "教学主文档（必读）：逐层讲清「为什么这样做」——六层架构、数据库四张表设计逻辑、双模式运行原理、六大服务模块底层算法（含转移矩阵/格网叠加/质心距离判定）、前端组件通信、MapLibre 渲染原理、部署与学习路径，并附实践题。",
    "外部工作清单.md": "项目文件夹之外需要完成的工作：PostgreSQL/PostGIS 安装验证、pgAdmin、Git、Docker、底图替换、县域 SHP 下载与坐标系转换教程，以及「外部→内部」四周对接计划。",
    "项目手册与学习指南.docx": "原始需求文档（教师下发）：项目需求与七个学习模块的原始定义，本仓库按此文档构建。",
    "启动项目.bat": "一键启动脚本（双击运行，界面为纯英文以规避中文代码页下 .bat 乱码解析问题）：先检测 8000/5173 端口是否已被占用（已占用则跳过启动），再分别打开后端（uvicorn）与前端（Vite --strictPort，端口被占会报错而不是偷偷换端口）两个命令行窗口；用 curl 轮询等待前端就绪（最多 30 秒）后自动打开浏览器 http://localhost:5173，并在窗口内给出排查提示。",
    "停止项目.bat": "停止脚本（双击运行，界面为纯英文）：关闭标题为 LandVISION-Backend / LandVISION-Frontend 的服务窗口，并用 netstat+taskkill 清理仍占用 8000/5173 端口的残留进程，便于干净地重新启动。",
    "中国_县.geojson": "全国县域行政区划原始数据（用户从天地图下载，gb 码= '156'+6 位标准码），是 tools/build_regions.py 的输入资产；体积较大，重新构建省市县数据时需要保留。",
    ".gitignore": "Git 忽略规则：排除 venv/、node_modules/、dist/、日志、缓存与 .env 等不入库内容。",
    ".github/workflows/ci.yml": "GitHub Actions 持续集成：提交/PR 时自动执行后端 pytest + 前端 npm run build，保证每次提交可运行。",
    ".vscode/settings.json": "VS Code 编辑器推荐配置（Python/Vue 插件建议、保存格式化等）。",

    # ---------- database/ ----------
    "database/00_create_database.sql": "建库脚本：创建 landvision 数据库并启用 PostGIS 扩展（执行顺序第 1 步）。",
    "database/01_init_schema.sql": "初始化表结构：parcels/pois/planning_control/regions 四张业务表的 DDL + 全部几何字段 GiST 空间索引，文件末尾附三条核心空间查询示例（视野查询/可达性/叠加分析）。",
    "database/02_seed_data.sql": "业务演示数据（v2.0）：10 宗地块（覆盖 10 类用地）、20 个 POI（交通/商业/教育/医疗/休闲）、3 条三区三线控制线（标准英文类型代码）；由 tools/generate_seed.py 生成，勿手改（执行顺序第 4 步）。",
    "database/03_regions.sql": "省级行政区数据：全国 34 个省级行政区（含边界几何）入库脚本，由 tools/generate_seed.py 生成（执行顺序第 3 步）。",
    "database/05_cities_counties.json": "全国市/县级行政区数据（346 市 + 2883 县，含边界），tools/build_regions.py 的产物，供 tools/load_regions_pg.py 写入 PostgreSQL。",

    # ---------- backend/ ----------
    "backend/.env": "本地环境变量（含数据库连接串 LANDVISION_DB_URL，含密码，不提交 Git）。",
    "backend/.env.example": "环境变量模板：复制为 .env 后填写数据库地址/密钥等。",
    "backend/alembic.ini": "Alembic 数据库迁移配置：指向 migrations/ 目录与 alembic_version 版本表。",
    "backend/requirements.txt": "Python 依赖清单（版本与 venv 实测一致）：fastapi/uvicorn/sqlalchemy/geoalchemy2/psycopg2/pydantic/shapely/alembic/pyshp/python-multipart + 测试依赖。",
    "backend/run.py": "一键启动脚本：等价于 uvicorn app.main:app --reload --port 8000。",
    "backend/Dockerfile": "后端生产镜像：安装依赖并暴露 8000 端口（配合 deploy/docker-compose.yml）。",
    "backend/__init__.py": "backend 包标记文件（空文件）。",
    "backend/examples/parcel_crud_practice.py": "教学练习脚本：用最小代码演示 FastAPI 路由 + Pydantic 校验 + 内存数据库三件套，独立运行于 8001 端口，与正式分层结构对比学习。",
    "backend/tests/test_smoke.py": "全链路冒烟测试（20 项，v3.0）：健康检查、地块 CRUD/期次筛选/锁定/批量删除/框选删除（delete-by-geometry）、SHP 导入端到端（内存 zip，含项目强制关联 422/404 校验）、POI 点要素导入（点面分离）、行政区、POI、三区三线体检（规则矩阵/判定依据/锁定）、项目 CRUD 与范围校验、四大结果持久化、地图标注、驾驶舱严格范围聚合（子范围不泄漏）、适宜性矛盾端点、综合分析报告（校验六章节与 12 大类）。",
    "backend/migrations/env.py": "Alembic 迁移运行环境：加载 app.models 的元数据，支持在线/离线执行迁移。",
    "backend/migrations/script.py.mako": "迁移脚本生成模板（alembic revision 命令用）。",
    "backend/migrations/versions/0001_initial.py": "迁移 0001：初始建表（parcels/pois/planning_control/change_records/regions + 空间索引）。",
    "backend/migrations/versions/0002_parcel_period.py": "迁移 0002：parcels 增加 period 期次字段与索引（支撑两期用地转移矩阵）。",
    "backend/migrations/versions/0003_drop_change_records.py": "迁移 0003：删除 change_records 表（变化监测模块下线后的数据库清理）。",
    "backend/migrations/versions/0004_v2_projects_and_results.py": "迁移 0004（v2.0 数据基础）：analysis_projects 分析项目表 + 四大结果表（land_change_patches/suitability_grids/planning_check_results/accessibility_results）+ map_features 标注表；parcels.period 非空默认基期；地块/POI/控制线增加项目归属与锁定字段。",
    "backend/migrations/versions/0005_standardize_zone_types.py": "迁移 0005（v2.0 术语规范）：planning_control.zone_type 由中文迁移为标准英文代码（permanent_basic_farmland/ecological_red_line/urban_growth_boundary），非标准类型记录移除。",
    "backend/migrations/versions/0006_patch_geom_generic.py": "迁移 0006（v2.0）：land_change_patches.geom 类型放宽为 GEOMETRY（消失/新增图斑可能为 MultiPolygon）。",
    "backend/migrations/versions/0007_gist_indexes_v3.py": "迁移 0007（v3.0）：land_change_patches/suitability_grids/map_features 补充 GiST 空间索引，范围过滤由 Seq Scan 升级为 Index Scan。",
    "backend/app/config.py": "配置中心（pydantic-settings）：应用名/版本/数据库地址/Demo 开关/CORS/上传上限/限流阈值等，环境变量前缀 LANDVISION_。",
    "backend/app/database.py": "数据访问层：SQLAlchemy 引擎懒初始化（连不上自动降级 Demo）、会话工厂、get_db 请求级会话依赖。",
    "backend/app/models.py": "ORM 模型（12 张表，v2.0）：AnalysisProject 分析项目、Parcel（period 必填/锁定/项目归属）、Poi、PlanningZone（标准三线英文代码）、Region 五张业务表 + land_change_patches/suitability_grids/planning_check_results/accessibility_results 四张结果表 + map_features 标注表；空间字段用 GeoAlchemy2 Geometry 声明。",
    "backend/app/schemas.py": "Pydantic 模型（v3.0）：12 大类用地枚举、标准三线代码与中文标签映射（ZONE_TYPE_LABELS）、地块/POI/控制线（含期次/项目/锁定）、分析项目、规则矩阵、批量操作、GeometryDelete 框选删除请求体与地图标注等请求响应模型、分页结构。",
    "backend/app/demo_data.py": "Demo 内存数据集：10 地块 + 20 POI + 3 条标准三线控制线 + 34 省级行政区（含几何），v2.0 内存持久化容器（PROJECTS/LAND_CHANGE_PATCHES/SUITABILITY_GRIDS/PLANNING_CHECK_RESULTS/ACCESSIBILITY_RESULTS/MAP_FEATURES）与 GeoJSON 便捷函数；由 tools/generate_seed.py 生成。",
    "backend/app/errors.py": "统一错误码与全局异常处理器：NOT_FOUND/CONFLICT/VALIDATION_ERROR/RATE_LIMITED/INTERNAL_ERROR 等，响应统一 {code, message, detail}。",
    "backend/app/middleware.py": "四件套中间件：X-Request-ID、安全响应头、IP 滑动窗口限流（120 次/分）、操作审计（写 logs/landvision.log）。",
    "backend/app/main.py": "应用入口：启动时探测数据库决定 Demo/POSTGIS 模式、注册中间件与 CORS、挂载 10 组路由、提供 / 与 /healthz 端点。",
    "backend/app/data/china_regions.json": "全国省市县行政区数据（Demo 模式懒加载；PostGIS 模式下由数据库提供），tools/build_regions.py 产物。",
    "backend/app/data/planning_rules.json": "体检规则矩阵配置（v2.0）：默认规则参考国土空间规划管控逻辑（耕地占基本农田=允许、建设用地占红线=冲突等），可通过 PUT /api/planning/rules 更新。",
    "backend/app/routers/__init__.py": "路由包说明。",
    "backend/app/routers/parcels.py": "地块路由（v3.0）：分页列表/GeoJSON（均支持 period 期次过滤）/CRUD/SHP 批量导入（强制关联项目与期次，缺失 422、项目不存在 404）/批量删除（跳过锁定）/按几何范围批量删除 delete-by-geometry（框选删除，intersects/within）/锁定解锁/批量设置期次。",
    "backend/app/routers/pois.py": "兴趣点路由（v3.0）：分页列表/GeoJSON/按类型过滤/新建/删除/批量删除（跳过锁定）/SHP 点要素导入（import，仅点要素，强制关联项目，点面分离）。",
    "backend/app/routers/analysis.py": "空间分析路由（模块一~三，v3.0）：转移矩阵（期次 SHP 导入/演示基期/计算，结果持久化，导入强制项目关联）+ 图斑查询、适宜性评价（刚性约束+格网持久化）+ 格网查询、可达性分析（结果持久化）+ 结果查询、盲区∩适宜区设施选址、适宜性矛盾提示（suitability/conflicts）、SHP 范围解析。",
    "backend/app/routers/planning.py": "三区三线体检路由（模块四，v3.0）：标准三线控制线管理（含批量删除/锁定/SHP 导入，导入强制项目关联）、规则矩阵查看与更新（GET/PUT rules）、单地块/任意几何体检（含判定依据）、批量体检（结果持久化）、变化图斑体检（模块一→四联动）、结果查询、台账 CSV 导出。",
    "backend/app/routers/projects.py": "分析项目路由（v2.0）：项目 CRUD —— 新建（名称+基期/末期年份+可选范围）、更新（范围变更需 confirm_scope_change 确认）、删除（结果级联删除）。",
    "backend/app/routers/map_features.py": "地图标注路由（v2.0）：地图绘制的点/线/面持久化（list/geojson/create/delete/batch-delete/lock）。",
    "backend/app/routers/dashboard.py": "数据驾驶舱路由（项目工作台）：POST /dashboard/summary 按项目与范围聚合全模块统计（持久化结果优先 + 流程进度/问题清单/规划建议，与报告生成同源）。",
    "backend/app/routers/report.py": "报告路由：生成综合分析报告（接收 project_id/scope 继承驾驶舱上下文，六章节结构化 JSON）/ 获取最近报告 / 下载 Markdown。",
    "backend/app/routers/regions.py": "行政区划路由：省/市/县列表与 GeoJSON/详情/下级/定位（中心+bbox）/SHP 导入。",
    "backend/app/routers/system.py": "系统路由：服务运行信息 / 最近操作审计日志。",
    "backend/app/services/__init__.py": "服务包说明。",
    "backend/app/services/spatial.py": "空间查询与地块 CRUD 服务（v3.0）：分页+bbox 视野查询（含 period 过滤）、GeoJSON 序列化、面积自动计算、编号唯一校验、锁定/批量删除/批量设置期次/按几何范围批量删除（delete_parcels_by_geometry，框选删除跳过锁定）；双模式分支的样板。",
    "backend/app/services/regions.py": "行政区服务：省/市/县查询、GeoJSON、定位（质心+包围盒）、SHP 导入。",
    "backend/app/services/shp_import.py": "SHP 导入服务（v3.0）：zip 解析（pyshp）→ 坐标系校验（仅 WGS84）→ 字段自动/手动映射 → 地类容错映射到 12 大类 → 批量入库；地块仅接受面要素、POI 仅接受点要素（import_pois_from_zip，点面分离）；require_project 强制项目关联（缺失 422、不存在 404），期次随导入直接落库。",
    "backend/app/services/planning_check.py": "三区三线体检服务（模块四，v3.0 规则矩阵版）：标准三线控制线管理（SHP 导入强制项目关联）、单地块/任意几何体检（规则矩阵判定 + 判定依据含亩数）、批量占用叠加 review_occupancy（结果持久化 planning_check_results）、变化图斑体检 review_patches（回写图斑冲突标记）、台账 CSV。",
    "backend/app/services/analysis.py": "空间分析服务（模块一~三，v3.0）：transition_matrix 两期叠加求交（结果持久化 land_change_patches）、期次导入（项目+期次随导入落库）、演示基期、suitability_evaluate 格网法多因子加权叠加（三区三线刚性约束强制不适宜 + 格网持久化）、accessibility_analyze 生活圈覆盖判定（结果持久化）、facility_sites 盲区∩适宜区选址、suitability_conflicts 适宜性矛盾提示（适宜格网∩体检冲突地块）；list_patches/list_grids 支持 scope 严格过滤；DEMO/POSTGIS 共用同一 shapely 实现。",
    "backend/app/services/planning_rules.py": "体检规则矩阵服务（v2.0，可配置）：加载/校验/写回 planning_rules.json（12 用地大类 × 三区三线 → conflict/warning/pass），提供 verdict_for/verdict_level/rules_table。",
    "backend/app/services/projects.py": "分析项目服务（v2.0）：项目 CRUD、范围 GeoJSON 合法性校验、范围变更确认机制、scope_within 子集校验与 resolve_project_scope 范围统一解析（各分析模块继承项目范围）。",
    "backend/app/services/map_features.py": "地图标注服务（v2.0）：地图绘制的点/线/面保存至 map_features 表，支持按项目过滤、锁定（锁定后不可删除）与批量删除。",
    "backend/app/services/report_gen.py": "驾驶舱统筹汇总 + 综合分析报告服务（v3.0，共用数据源）：collect_dashboard 优先从持久化结果表读取（land_change_patches/suitability_grids/planning_check_results/accessibility_results），无结果时实时计算兜底；v3.0 严格范围聚合 —— 全部数据源按 scope 过滤，跨界要素面积按交集裁剪（POSTGIS 用 ST_Intersection，Demo 用 shapely+STRtree），可达性子范围只读实时复算；组装项目概况、流程进度、问题清单与规划建议；to_markdown 输出六章节报告。",

    # ---------- frontend/ ----------
    "frontend/package.json": "前端依赖与脚本：vue/vite/pinia/vue-router/axios/maplibre-gl/echarts/element-plus；dev/build/preview 命令。",
    "frontend/package-lock.json": "依赖版本锁定文件（npm install 精确复现）。",
    "frontend/vite.config.js": "Vite 配置：开发服务器 5173 端口、/api 反向代理到后端 8000、构建输出 dist/。",
    "frontend/index.html": "SPA 入口 HTML：挂载 #app 节点并加载 src/main.js。",
    "frontend/.env.development": "前端开发环境变量（Vite 读取，如 API 地址前缀）。",
    "frontend/nginx.conf": "生产 Nginx 配置：托管构建产物、/api 反向代理到后端容器、history 路由回退到 index.html。",
    "frontend/Dockerfile": "前端生产镜像：多阶段构建（node 构建 → nginx 托管），暴露 80 端口。",
    "frontend/src/main.js": "应用入口：创建 Vue 应用、注册 Pinia 与 Router、全局注册全部 Element Plus 图标、应用主题。",
    "frontend/src/App.vue": "主布局与全局导航（v2.0）：折叠侧边栏（菜单分组 + 推荐流程编号 ①~⑥）、顶栏（分析项目选择/新建入口、搜索/通知/主题/用户/帮助）、路由出口。",
    "frontend/src/router/index.js": "Vue Router：7 条路由（Dashboard/地块管理/转移矩阵/适宜性/可达性/三区三线体检/报告），路由切换时同步 document.title。",
    "frontend/src/stores/parcel.js": "Pinia 地块状态：地块列表/地块 GeoJSON/POI GeoJSON/规划要素 GeoJSON、加载 actions、总面积 getter；各页面共享一处更新处处联动。",
    "frontend/src/stores/ui.js": "Pinia UI 状态（v3.0）：深浅主题、图层可见性、行政区筛选、框选结果、分析项目上下文（currentProject/currentProjectId）、模块联动参数（linkedPatches/linkedScope/linkedFacilityTypes）、分析版本号 analysisVersion（各分析页成功后自增，驾驶舱自动刷新）、顶栏通知。",
    "frontend/src/api/index.js": "接口封装汇总（v3.0）：按后端 10 组路由分组导出全部 API 函数（含 multipart 上传与 blob 下载），含 POI 导入 importPoisShp、框选删除 deleteParcelsByGeometry、适宜性矛盾 getSuitabilityConflicts；页面不直接写 URL。",
    "frontend/src/utils/request.js": "Axios 实例封装：baseURL=/api、请求/响应拦截器、统一错误提示（适配后端 {code,message,detail}）、超时设置。",
    "frontend/src/utils/colors.js": "配色规范（v3.0）：12 大类国标色、POI 类型色、变化图斑色；三区三线标准术语与制图样式（ZONE_TYPE_LABELS/ZONE_TYPE_COLORS/ZONE_TYPE_LINE_STYLES：红线 #E53935 实线 3px、基本农田 #FFB300 实线 2.5px、开发边界 #1E88E5 虚线 2px）；体检结论配色；图例元数据（ZONE_LEGEND 附线型 line 元数据，图例面板预览实线/虚线样式）。",
    "frontend/src/utils/geo.js": "前端几何工具：坐标距离（米）、多边形面积（㎡）、点/要素与框选/圈选/多边形的包含判断（地图框选统计的底层实现）。",
    "frontend/src/directives/drag.js": "v-drag 自定义指令：让地图上的图标栏/浮动面板可拖拽移动。",
    "frontend/src/styles/main.css": "全局样式与设计令牌：CSS 变量（主题色/圆角/阴影/文本色）、玻璃拟态面板、页面图标按钮、通用布局类。",
    "frontend/src/components/MapView.vue": "核心地图组件（v3.0）：OpenFreeMap 底图+卫星影像、图层管理（地块/POI/三区三线规范线型/变化图斑/适宜性格网/可达性覆盖/行政区边界）、框选圈选多边形选择统计（可配置「删除选中地块」按钮 selectionDelete/selection-delete）、距离面积测量、悬停高亮、点击弹窗、批量选择模式（抛 batch-selection）、保存绘制（抛 save-drawing）、锁定地块虚线描边、图例线型预览（三区三线实线/虚线+线宽）、地图视野持久化、对外暴露 flyTo/fitBounds/clearBatchSelection/clearDrawing/clearSelection/getMap。",
    "frontend/src/components/RegionSelector.vue": "行政区选择器：国家→省→市→县四级级联检索、定位飞行、点击外部关闭；固定在地图右上角。",
    "frontend/src/components/StatsCard.vue": "专业指标卡：数值/单位/图标/分组标签/环比变化/副文本/迷你趋势折线。",
    "frontend/src/components/ParcelInfo.vue": "地块详情抽屉（v2.0）：地块属性（期次/所属项目/锁定状态）+ 三区三线体检结论（重叠亩数 + 冲突/警告判定依据）+ 地图定位 + 模块跳转（查看体检结果/查看转移矩阵记录，自动携带地块 ID）。",
    "frontend/src/views/DashboardView.vue": "数据驾驶舱（v3.0 项目工作台）：当前项目信息（名称/基期末期年份/范围，含「已按此范围聚合」严格范围徽标）与流程进度条（①~⑥ 完成状态）、范围划定/变更对话框（行政区任意层级/SHP，变更确认机制）、关键结论摘要卡片（冲突地块/设施盲区/耕地净减少/高度适宜占比，点击下钻对应模块）、问题识别与规划建议清单（违规变化→转移矩阵页钻取）、用地结构/体检占用/适宜性/转移矩阵图表；数据来自 /api/dashboard/summary（持久化结果优先+严格范围聚合），watch analysisVersion 自动刷新。",
    "frontend/src/views/ParcelManagementView.vue": "地块管理（v3.0）：期次筛选（全部/基期/末期，表格与地图同步）、上传强制关联项目与期次（前端先行校验）、地图批量选择模式（点击地块/POI/控制线进入选中集 → 批量删除/锁定）、框选删除模式（地图框选/圈选范围 → delete-by-geometry 批量删除，锁定项跳过）、兴趣点管理面板（POI 点要素 SHP 导入/列表/删除）、锁定地块（虚线描边，锁定后不可删除）、标注面板（地图绘制保存至 map_features 并支持锁定/删除）、详情抽屉（期次/项目/判定依据/模块跳转）。",
    "frontend/src/views/TransferMatrixView.vue": "模块一 用地变化转移矩阵（v3.0）：基期/末期 SHP 分次导入（强制关联项目）、一键生成演示基期、范围继承项目、矩阵透视表（色深=面积）、消失/新增图斑标签、面积增减图、变化图斑上图；模块联动按钮：对变化图斑进行合规检查（→体检）、评估新增用地适宜性（→适宜性）、分析新增用地设施可达性（→可达性，图斑并集范围随 ui store 传递，并按矩阵结果预置设施类型）。",
    "frontend/src/views/SuitabilityView.vue": "模块二 土地适宜性评价（v3.0）：评价目标切换（建设/耕地）、因子权重滑杆（自动归一化）、范围默认继承项目（转移矩阵联动范围优先）、40×40 格网四级适宜性专题图（永久基本农田/生态保护红线刚性约束强制不适宜）、等级分布图、格网点击弹分、适宜性矛盾提示（高度/中等适宜 ∩ 体检冲突地块，评价完成后自动校验）、「查看体检结论」反向校验联动。",
    "frontend/src/views/AccessibilityView.vue": "模块三 服务设施可达性分析（v3.0）：设施类型多选（接收转移矩阵联动预置并提示）、服务半径滑杆+预设、范围默认继承项目（联动范围优先）、覆盖率饼图、盲区清单（行点击定位）、覆盖/盲区红绿专题图与图斑弹窗、「推荐设施选址」（盲区∩适宜布局区，蓝色虚线展示候选区域）。",
    "frontend/src/views/PlanningCheckView.vue": "模块四 三区三线合规性批量体检（v3.0）：标准三线控制线（SHP 导入选择约束类型并强制关联项目/图上绘制/批量删除/锁定）、规则矩阵对话框（12 地类 × 三线）、批量体检（结果持久化，台账含判定依据与重叠亩数）、变化图斑联动体检（接收转移矩阵参数自动执行）、判定依据明细表、台账 CSV 导出（含判定依据）。",
    "frontend/src/views/ReportView.vue": "报告生成（v3.0 综合分析）：继承驾驶舱分析项目与范围（含「已按此范围聚合」徽标）；六章节预览（项目概况含流程进度/现状评价/问题识别/原因分析含判定依据/规划建议/附录数据表）；问题与台账行可点击跳转对应模块/地块（钻取，违规变化→转移矩阵）；下载 Markdown。",

    # ---------- tools/ ----------
    "tools/generate_seed.py": "种子数据生成器（单一数据源）：生成 backend/app/demo_data.py、database/02_seed_data.sql、database/03_regions.sql 三份数据文件，保证 Demo 与数据库数据一致。",
    "tools/build_regions.py": "行政区数据构建：读取 china_provinces.json 与 中国_县.geojson，在线补充市/县边界，产出 backend/app/data/china_regions.json 与 database/05_cities_counties.json。",
    "tools/load_regions_pg.py": "行政区入库脚本：把 database/05_cities_counties.json 的省市县数据（含边界）写入 PostgreSQL regions 表；用法 python tools/load_regions_pg.py <用户> <密码>。",
    "tools/reset_pg_password.bat": "Windows 批处理：重置 PostgreSQL 用户密码为 postgres（ALTER USER 语句）。",
    "tools/generate_manual_docx.py": "本说明书生成器：遍历项目文件、注入逐文件说明、排版输出《项目说明书.docx》。",
    "tools/v3_postgis_live_test.ps1": "v3.0 POSTGIS 实测脚本（PowerShell，UTF-8 BOM）：对运行中的后端逐项断言 POI 导入校验（422/404）、上传强制项目校验（四入口）、点要素 POI 导入、框选删除（锁定跳过）、可达性/适宜性/体检持久化、驾驶舱严格范围聚合（子范围收敛为 0）、适宜性矛盾端点，并自动清理测试项目/POI/地块（无 BOM JSON 文件 + curl --data-binary）。",
    "tools/assets/china_provinces.json": "中国省级行政区 GeoJSON（DataV 公开数据，含边界），build_regions.py 的输入。",
    "tools/assets/province_children/*.json": "各省下辖市县边界数据缓存（34 个，如 420000.json=湖北省）：build_regions.py 在线下载的本地缓存，避免重复请求。",

    # ---------- deploy/ ----------
    "deploy/docker-compose.yml": "生产编排：postgis（数据库，挂载建表脚本）/ backend（FastAPI，强制 POSTGIS 模式）/ frontend（Nginx）三服务一键启动，访问 http://localhost:8080。",

    # ---------- docs/ ----------
    "docs/01-项目概述.md": "项目定位、业务场景、功能清单（四分析模块）、双模式运行、技术栈总览。",
    "docs/02-技术架构.md": "六层架构图、一次请求的完整链路（以可达性分析为例）、关键设计决策、Mermaid 数据流全景图。",
    "docs/03-目录结构.md": "逐文件职责索引（根目录/database/backend/frontend/deploy/docs）。",
    "docs/04-数据库设计.md": "连接配置、业务表与 v2.0 项目/结果/标注表字段详解（含 parcels.period 必填默认基期、标准三线代码、锁定字段）、12 大类用地枚举、GiST 空间索引、核心 SQL 速查、脚本执行顺序。",
    "docs/05-后端接口设计.md": "统一约定（错误码/分页/安全运维）+ 10 组路由接口表（parcels/regions/pois/planning 体检/analysis 三模块/projects 项目/map-features 标注/dashboard 统筹/report 综合分析/system）。",
    "docs/06-前端页面设计.md": "技术栈、地图优先设计原则、目录职责、MapView 能力清单（含 cells/coverage 叠加层）、RegionSelector、7 页面明细、设计规范。",
    "docs/07-启动调试指南.md": "后端（Demo/POSTGIS 两种模式）与前端启动步骤、数据库接入、常见调试方法。",
    "docs/08-常见问题排查.md": "常见问题与解决方案（端口占用/依赖安装/坐标系/中文编码/地图加载等）。",
    "docs/09-企业级升级说明.md": "企业级改造清单（错误码/分页/审计/限流/健康检查/Alembic/CI）、v1.2~v1.3 业务调整记录、验证方式、剩余路线图。",
}

# ---------------------------------------------------------------------------
# 文档结构
# ---------------------------------------------------------------------------

MODULE_TABLE = [
    ("分组", "页面", "后端", "核心能力"),
    ("总览", "数据驾驶舱（项目工作台）", "/api/dashboard", "项目信息 + 流程进度 + 关键结论摘要下钻 + 问题识别/规划建议 + 项目范围管理（持久化结果优先取数）"),
    ("数据管理", "地块管理", "/api/parcels", "分页/CRUD/视野查询/期次筛选/SHP 导入（强制关联项目期次）/批量删除锁定/框选删除 delete-by-geometry/POI 点要素管理/地图标注"),
    ("上下文", "分析项目（顶栏）", "/api/projects", "项目 CRUD、范围统一继承与变更确认、子集校验；各模块自动继承范围与期次"),
    ("模块一", "用地转移矩阵", "/api/analysis/transition", "两期叠加求交 → 转移矩阵 + 图斑持久化（land_change_patches）+ 联动体检/适宜性/可达性"),
    ("模块二", "适宜性评价", "/api/analysis/suitability", "格网法多因子加权叠加 + 三区三线刚性约束 + 格网持久化 + 反向校验体检（适宜×冲突矛盾提示）"),
    ("模块三", "设施可达性", "/api/analysis/accessibility", "生活圈覆盖 + 结果持久化 + 盲区∩适宜区设施选址建议"),
    ("模块四", "三区三线体检", "/api/planning", "标准三线 + 可配置规则矩阵（判定依据）+ 结果持久化 + 图斑体检 + 台账 CSV 导出"),
    ("输出", "报告生成", "/api/report", "继承项目与范围，综合分析六章节 + 问题清单/规划建议 + Markdown 下载"),
]

API_TABLE = [
    ("路由前缀", "功能", "关键接口"),
    ("/api/parcels", "地块管理", "分页/GeoJSON（period 过滤）/ CRUD / import-shp（项目+期次强制）/ batch-delete / delete-by-geometry（框选删除）/ {id}/lock / batch-set-period"),
    ("/api/regions", "行政区划", "省/市/县列表 / GeoJSON / 详情 / 下级 / locate 定位 / import"),
    ("/api/pois", "兴趣点", "分页列表 / GeoJSON / 新建 / 删除 / batch-delete / import（点要素 SHP，强制项目关联）"),
    ("/api/analysis", "空间分析（模块一~三）", "transition（import-shp/generate-demo-base/matrix/patches）、suitability（targets/evaluate/grids/conflicts 矛盾提示）、accessibility（analyze/results）、facility-sites、parse-scope"),
    ("/api/planning", "三区三线体检（模块四）", "zones（含 import-shp/batch-delete/lock）/ rules（规则矩阵查看与更新）/ check / review / review-patches / results / review/export"),
    ("/api/projects", "分析项目", "GET/POST /projects、GET/PUT/DELETE /projects/{id}（范围变更需确认）"),
    ("/api/map-features", "地图标注", "list / geojson / create / delete / batch-delete / {id}/lock"),
    ("/api/dashboard", "数据驾驶舱（统筹）", "summary —— 按项目与范围聚合（持久化结果优先 + 流程进度/问题/建议）"),
    ("/api/report", "报告生成（综合分析）", "generate（含 project_id）/ latest / latest/download"),
    ("/api/system", "系统运维", "info / audit（另有 /healthz 健康检查）"),
]

DB_TABLE = [
    ("表", "空间类型", "核心字段", "业务角色"),
    ("parcels 地块", "Polygon", "parcel_code/name/land_use(12大类)/district/region_code/area_sqm/period(必填)/project_id/locked", "平台主角；期次区分基期/末期；锁定后不可删除"),
    ("pois 兴趣点", "Point", "name/poi_type/project_id/locked", "地图展示、可达性分析的设施数据源"),
    ("planning_control 三区三线", "Polygon", "zone_name/zone_type(标准三线代码)/zone_level/project_id/locked", "三区三线体检的刚性约束（用户自行导入）"),
    ("regions 行政区划", "MultiPolygon", "code/name/level/parent_code", "省→市→县下钻、定位、分析范围划定"),
    ("analysis_projects 分析项目", "-", "name/base_year/current_year/scope_geojson", "分析业务上下文：范围与期次统一继承"),
    ("land_change_patches 变化图斑", "GEOMETRY", "project_id/from_land_use/to_land_use/change_type/is_conflict", "转移矩阵结果持久化，模块间复用"),
    ("suitability_grids 适宜性格网", "Polygon", "project_id/score/level/factors_json", "适宜性结果持久化"),
    ("planning_check_results 体检结果", "-", "project_id/parcel_id/zone_id/overlap_area_sqm/conclusion", "体检结果持久化，可追溯/复用"),
    ("accessibility_results 可达性结果", "-", "project_id/facility_types/radius_m/coverage_rate/gap_parcel_ids", "可达性结果持久化"),
    ("map_features 地图标注", "GEOMETRY", "project_id/name/feature_type/locked", "地图绘制的点/线/面持久化"),
]


# ---------------------------------------------------------------------------
# Word 排版辅助
# ---------------------------------------------------------------------------

def set_font(run, name_cn="宋体", size=10.5, bold=False, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name_cn)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def heading(doc, text, level=1):
    h = doc.add_heading("", level=level)
    run = h.add_run(text)
    set_font(run, "黑体", {1: 16, 2: 13, 3: 11.5}[level], bold=True,
             color=(0x1F, 0x3B, 0x5C) if level <= 2 else (0x2E, 0x86, 0xAB))
    h.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    h.paragraph_format.space_after = Pt(6)
    return h


def para(doc, text, size=10.5, bold=False, indent=0, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(indent)
    run = p.add_run(text)
    set_font(run, "宋体", size, bold, color)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    set_font(run, "宋体", 10.5)
    return p


def table(doc, rows, widths=None):
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            cell = t.cell(i, j)
            cell.text = ""
            run = cell.paragraphs[0].add_run(cell_text)
            set_font(run, "宋体", 9.5, bold=(i == 0))
    return t


def file_entry(doc, path, desc):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(path)
    set_font(r1, "等线", 10, bold=True, color=(0x2E, 0x86, 0xAB))
    r2 = p.add_run(" —— " + desc)
    set_font(r2, "宋体", 10)
    return p


def walk_files():
    out, missing = [], []
    for dirpath, dirnames, filenames in os.walk(ROOT):
        dirnames[:] = [d for d in dirnames if d not in EXCLUDE_DIRS]
        rel = os.path.relpath(dirpath, ROOT).replace("\\", "/")
        for fn in sorted(filenames):
            if fn in EXCLUDE_FILES or fn.startswith("~$"):
                continue
            path = f"{rel}/{fn}" if rel != "." else fn
            if "province_children/" in path:
                path = "tools/assets/province_children/*.json"
            out.append(path)
            if path not in DESCRIPTIONS:
                missing.append(path)
    seen, uniq = set(), []
    for p in sorted(out):
        if p not in seen:
            seen.add(p)
            uniq.append(p)
    return uniq, missing


# ---------------------------------------------------------------------------
# 主流程
# ---------------------------------------------------------------------------

def build():
    doc = Document()
    # 页面默认样式
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(10.5)

    # ---------- 封面 ----------
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("LandVISION")
    set_font(run, "黑体", 36, bold=True, color=(0x2E, 0x86, 0xAB))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("国土空间数据管理与智能分析可视化平台")
    set_font(run, "黑体", 22, bold=True, color=(0x1F, 0x3B, 0x5C))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("项 目 说 明 书")
    set_font(run, "黑体", 28, bold=True)
    for _ in range(2):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("细化到每个文件的职责说明 · v3.0")
    set_font(run, "宋体", 13, color=(0x66, 0x6E, 0x79))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("PostgreSQL/PostGIS · FastAPI · Vue 3 + MapLibre GL JS · ECharts · Element Plus")
    set_font(run, "宋体", 12, color=(0x66, 0x6E, 0x79))
    doc.add_page_break()

    # ---------- 一、项目概述 ----------
    heading(doc, "一、项目概述", 1)
    para(doc, "LandVISION 是一个面向国土空间规划与自然资源管理场景的全栈 WebGIS 教学实战项目，覆盖数据管理、"
              "空间分析与可视化报告全流程。系统采用双模式运行：默认连接本机 PostgreSQL + PostGIS（POSTGIS 模式）；"
              "数据库不可用时自动降级为内存演示数据（DEMO 模式），保证任何环境下都能跑通全链路。", indent=0)
    heading(doc, "1.1 技术栈", 2)
    bullet(doc, "后端：Python 3.14 · FastAPI · SQLAlchemy 2.0 · GeoAlchemy2 · Shapely · Alembic · pyshp")
    bullet(doc, "前端：Vue 3 · Vite · Pinia · Vue Router · Axios · MapLibre GL JS · ECharts · Element Plus")
    bullet(doc, "数据库：PostgreSQL 16 + PostGIS 3.x（parcels/pois/planning_control/regions 四张表 + GiST 空间索引）")
    bullet(doc, "部署：Docker Compose（数据库 + 后端 + Nginx 前端）+ GitHub Actions CI")
    heading(doc, "1.2 业务模块（7 个页面）", 2)
    table(doc, MODULE_TABLE)

    # ---------- 二、系统架构 ----------
    heading(doc, "二、系统架构", 1)
    para(doc, "系统分六层：① 数据存储层（PostgreSQL+PostGIS，五张业务表 + 项目/结果/标注等十张表）→ ② 数据访问层（SQLAlchemy+GeoAlchemy2）"
              "→ ③ 业务服务层（spatial/regions/shp_import/planning_check/analysis/report_gen/projects/map_features/planning_rules 九个服务模块）"
              "→ ④ 接口层（FastAPI 10 组路由）→ ⑤ 前端框架层（Vue 3 + Pinia + Router + Axios）"
              "→ ⑥ 可视化层（MapLibre 地图 + ECharts 图表 + Element Plus）。")
    para(doc, "关键设计决策：几何一律以 GeoJSON（EPSG:4326）传输；空间运算统一由服务层实现（DEMO 用 shapely、"
              "POSTGIS 用 ST_* SQL，结果一致便于教学对比）；前端地图与表格通过 Pinia 状态共享联动；"
              "拖动地图按视野 bbox 增量加载地块（GiST 索引两段式过滤）。")
    para(doc, "一次请求的完整链路（以模块三可达性分析为例）：页面点击分析 → Axios POST /api/analysis/accessibility/analyze"
              "→ FastAPI 校验参数 → services/analysis.accessibility_analyze 对地块质心与所选设施逐一求距 → "
              "≤ 服务半径判覆盖、否则进盲区清单 → 返回覆盖率/盲区/覆盖专题 GeoJSON → 前端渲染卡片 + 表格 + 红绿图层。")

    # ---------- 三、数据库设计 ----------
    heading(doc, "三、数据库设计", 1)
    table(doc, DB_TABLE)
    para(doc, "四张表全部几何字段均建 GiST 空间索引；用地性质遵循 GB/T 21010-2017 一级类（12 大类：耕地/园地/林地/"
              "草地/商服用地/工矿仓储用地/住宅用地/公共管理与公共服务用地/特殊用地/交通运输用地/水域及水利设施用地/其他土地）。", indent=0)
    para(doc, "初始化顺序：00_create_database.sql（建库+扩展）→ 01_init_schema.sql（建表+索引）→ 03_regions.sql"
              "（省级行政区）→ 02_seed_data.sql（业务演示数据）；后续结构变更走 Alembic 迁移（0001 初始 → 0002 期次字段 → 0003 删除变化监测表"
              "→ 0004 v2.0 项目与结果表 → 0005 三线标准代码 → 0006 图斑几何通用化 → 0007 v3.0 结果表 GiST 空间索引）。", indent=0)

    # ---------- 四、逐文件说明 ----------
    heading(doc, "四、逐文件说明（核心章节）", 1)
    para(doc, "以下按目录逐个说明每个文件的职责。文件路径为相对项目根目录的路径；venv/、node_modules/、dist/、"
              "缓存与日志等可再生成内容不列入。", size=10, color=(0x66, 0x6E, 0x79))

    files, missing = walk_files()
    groups = [
        ("4.1 根目录", lambda p: "/" not in p),
        ("4.2 database/ 数据库脚本", lambda p: p.startswith("database/")),
        ("4.3 backend/ 后端", lambda p: p.startswith("backend/")),
        ("4.4 frontend/ 前端", lambda p: p.startswith("frontend/")),
        ("4.5 tools/ 工具脚本与数据资产", lambda p: p.startswith("tools/")),
        ("4.6 deploy/ 部署", lambda p: p.startswith("deploy/")),
        ("4.7 docs/ 学习文档", lambda p: p.startswith("docs/")),
    ]
    covered = set()
    for title, pred in groups:
        heading(doc, title, 2)
        for path in files:
            if pred(path) and path not in covered:
                covered.add(path)
                file_entry(doc, path, DESCRIPTIONS[path])
    leftover = [p for p in files if p not in covered]
    if leftover:
        heading(doc, "4.8 其他", 2)
        for path in leftover:
            file_entry(doc, path, DESCRIPTIONS.get(path, "（无说明）"))

    # ---------- 五、快速开始 ----------
    heading(doc, "五、快速开始", 1)
    heading(doc, "5.1 启动后端（Demo 模式）", 2)
    para(doc, "cd backend  →  ..\\venv\\Scripts\\python.exe -m uvicorn app.main:app --reload --port 8000", indent=14)
    para(doc, "API 文档（Swagger UI）：http://127.0.0.1:8000/docs ；健康检查：GET /healthz。", indent=14)
    heading(doc, "5.2 启动前端", 2)
    para(doc, "cd frontend  →  npm install（首次）→  npm run dev", indent=14)
    para(doc, "访问 http://localhost:5173（Vite 将 /api 代理到后端 8000）。", indent=14)
    heading(doc, "5.3 接入 PostgreSQL（POSTGIS 模式）", 2)
    para(doc, "执行 database/00→01→03→02 脚本建库灌数；tools/load_regions_pg.py 导入省市县行政区；"
              "backend/.env 配置 LANDVISION_DB_URL 后重启后端，日志显示「POSTGIS」即为接入成功。", indent=14)
    heading(doc, "5.4 测试与构建", 2)
    para(doc, "后端测试：cd backend  →  ..\\venv\\Scripts\\python.exe -m pytest（20 项冒烟测试）", indent=14)
    para(doc, "v3.0 POSTGIS 实测：后端启动后执行 tools\\v3_postgis_live_test.ps1（23 项断言，自动清理测试数据）", indent=14)
    para(doc, "前端构建：cd frontend  →  npm run build", indent=14)
    para(doc, "生产部署：cd deploy  →  docker compose up -d --build（访问 http://localhost:8080）", indent=14)

    # ---------- 六、接口清单 ----------
    heading(doc, "六、后端接口清单（10 组路由）", 1)
    table(doc, API_TABLE)

    # ---------- 附录 ----------
    heading(doc, "附录 A：四大分析模块算法速览", 1)
    bullet(doc, "模块一 转移矩阵：两期地块叠加求交（base×current），按地类组合聚合转换面积；基期未保留部分=消失、末期新出现部分=新增（虚拟地类）。")
    bullet(doc, "模块二 适宜性评价：范围内 40×40 格网，逐格计算各因子得分（0~100），权重归一化后加权求和；≥80 高度适宜 / 60~80 中等 / 40~60 勉强 / <40 不适宜。")
    bullet(doc, "模块三 可达性分析：地块质心到所选设施直线距离 ≤ 服务半径（默认 800m，15 分钟生活圈口径）判覆盖；输出覆盖率与盲区清单。")
    bullet(doc, "模块四 三区三线体检：地块 × 审查要素叠加求交得占用面积矩阵；占用永久基本农田/生态保护红线=冲突、其余重叠=提示；台账支持 CSV 导出与图上定位。")
    heading(doc, "附录 B：说明", 1)
    bullet(doc, "本说明书由 tools/generate_manual_docx.py 自动生成；修改说明文字后重新运行该脚本即可更新。")
    bullet(doc, "变化监测模块（原遥感栅格差值检测）已按业务意见整体下线：monitoring 路由、change_detection 服务、"
              "监测页面、分屏对比组件、change_records 表（Alembic 0003 迁移）及相关文档均已删除。")

    doc.save(OUT_DOCX)
    return files, missing


if __name__ == "__main__":
    all_files, miss = build()
    print(f"[OK] 已生成：{OUT_DOCX}")
    print(f"[OK] 覆盖文件数：{len(all_files)}")
    if miss:
        print("[WARN] 以下文件缺少说明：")
        for m in miss:
            print("   ", m)
    else:
        print("[OK] 所有文件均有说明，无遗漏。")
